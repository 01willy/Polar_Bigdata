# -*- coding: utf-8 -*-
"""mk_qa_doc.py — 발표 대비 예상 질문 36건 답변 문서 (docx → soffice pdf).

입력: scratchpad/qa_answers.md (마크다운, ## Qn. 제목 / **라벨.** 본문 구조)
출력: deck/render/발표QA대비_ALT_ctrl.docx
실행: cd /home/willy010313/Polar_Bigdata && python3 <this file>
"""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

MD_IN = "deck/qa_answers.md"
OUT = "deck/render/발표QA대비_ALT_ctrl.docx"
FONT = "Pretendard"
ACC = (0xEA, 0x85, 0x1B)
INK = (0x11, 0x11, 0x11)
GRAY = (0x6B, 0x6B, 0x6B)
GRAY2 = (0x4A, 0x4A, 0x4A)


def set_font(run, size=10.5, bold=False, color=None, name=FONT):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def add_runs(p, text, size=10.5, color=INK, bold_color=INK):
    """**bold** 인라인 처리."""
    for i, seg in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not seg:
            continue
        if i % 2 == 1:
            set_font(p.add_run(seg), size, True, bold_color)
        else:
            set_font(p.add_run(seg), size, False, color)


doc = Document()
for sec in doc.sections:
    sec.top_margin = Mm(18); sec.bottom_margin = Mm(18)
    sec.left_margin = Mm(20); sec.right_margin = Mm(20)

# ---------- 표지부 ----------
p = doc.add_paragraph()
set_font(p.add_run("발표 질의응답 대비 문서"), 18, True)
p = doc.add_paragraph()
set_font(p.add_run("영구동토 활동층 두께(ALT) 예측 연구 · 예상 질문 36건과 답변"), 12, False, GRAY2)
p = doc.add_paragraph()
set_font(p.add_run("팀 ALT_ctrl · 2026 극지 빅데이터·인공지능 활용 경진대회 본선 · 2026. 08. 28"),
         10, False, GRAY)
p = doc.add_paragraph()
set_font(p.add_run("모든 수치는 예선 보고서(main.tex) 정합값이다. 페이지 번호는 본선 발표덱(23장) 기준이다. "
                   "각 답변은 한줄 답, 쉬운 설명(예시 포함), 핵심 수치, 근거의 순서로 구성한다."), 9.5, False, GRAY)

md = open(MD_IN, encoding="utf-8").read()
lines = md.split("\n")

LABELS = ("질문.", "요약.", "상세.", "근거.")
for raw in lines:
    line = raw.rstrip()
    if not line.strip():
        continue
    m = re.match(r"^##\s+(Q\d+)\.\s*(.*)$", line)
    if m:
        doc.add_paragraph()  # 간격
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(m.group(1) + ".  "), 13, True, ACC)
        set_font(p.add_run(m.group(2)), 13, True, INK)
        continue
    if re.match(r"^#\s", line):   # 그룹 대제목 등은 무시하거나 소제목 처리
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        set_font(p.add_run(line.lstrip("# ").strip()), 14, True, GRAY2)
        continue
    stripped = line.strip()
    # 라벨 단락 (**질문.** ... 등)
    lab = re.match(r"^\*\*(질문|한줄 답|쉽게 설명하면|숫자로|근거|심화 방어\(요약\))\.\*\*\s*(.*)$", stripped)
    if lab:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        name = lab.group(1)
        color = GRAY if name in ("질문", "근거") else INK
        accented = name in ("한줄 답", "쉽게 설명하면", "숫자로", "심화 방어(요약)")
        set_font(p.add_run(name + "  "), 10.5, True, ACC if accented else GRAY2)
        add_runs(p, lab.group(2), 10 if name in ("질문", "근거") else 10.5, color)
        continue
    # 불릿
    mb = re.match(r"^[-*]\s+(.*)$", stripped)
    if mb:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Mm(6)
        p.paragraph_format.space_after = Pt(1)
        set_font(p.add_run("·  "), 10.5, True, GRAY2)
        add_runs(p, mb.group(1), 10.5, INK)
        continue
    # 표 행(마크다운 표)은 고정 스타일 소형 텍스트
    if stripped.startswith("|"):
        if set(stripped.replace("|", "").replace("-", "").replace(":", "").strip()) == set():
            continue  # 구분선 스킵
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Mm(4)
        p.paragraph_format.space_after = Pt(0)
        cells = [c.strip() for c in stripped.strip("|").split("|")]
        add_runs(p, "   ".join(cells), 9.5, GRAY2)
        continue
    # 일반 단락
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_runs(p, stripped, 10.5, INK)

doc.save(OUT)
print("saved", OUT)
